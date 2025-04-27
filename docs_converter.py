"""
   CodeCraft PMS Backend Project

   파일명   : docs_converter.py                                                          
   생성자   : 김창환                                                          
                                                                               
   생성일   : 2024/11/26                                                      
   업데이트 : 2024/12/14                                                   
                                                                               
   설명     : DB로부터 정보를 받아와 산출물을 문서화 해주는 기능 정의
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, date
from urllib.parse import quote
from logger import logger
import pymysql, os, sys, traceback
import logging, requests
import re  # 정규식 사용

sys.path.append(os.path.abspath('/data/Database Project'))
import output_DB
import push

router = APIRouter()

class ConverterPayload(BaseModel):
    doc_type: int
    doc_s_no: int

def replace_placeholder_in_cell(cell, placeholder, replacement):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

def process_meeting_minutes(doc_s_no):
    """
    회의록 문서를 생성하는 기능을 처리합니다.
    """
    meeting_data = output_DB.fetch_one_meeting_minutes(doc_s_no)
    if not meeting_data:
        raise HTTPException(status_code=404, detail="Meeting data not found")

    try:
        doc = Document("/data/Docs_Template/회의록.docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template loading error: {e}")

    # 자리표시자 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_in_cell(cell, "{f1}", meeting_data.get("doc_m_title", ""))
                replace_placeholder_in_cell(
                    cell, "{f2}",
                    meeting_data.get("doc_m_date", "").strftime("%Y-%m-%d")
                    if isinstance(meeting_data.get("doc_m_date"), date)
                    else str(meeting_data.get("doc_m_date", ""))
                )
                replace_placeholder_in_cell(cell, "{f3}", meeting_data.get("doc_m_manager", ""))
                replace_placeholder_in_cell(cell, "{f4}", meeting_data.get("doc_m_loc", ""))
                replace_placeholder_in_cell(cell, "{f5}", str(len(meeting_data.get("doc_m_member", "").split(";"))))
                replace_placeholder_in_cell(cell, "{f6}", meeting_data.get("doc_m_content", ""))
                replace_placeholder_in_cell(cell, "{f7}", meeting_data.get("doc_m_result", ""))

                members = [member.split(",") for member in meeting_data.get("doc_m_member", "").split(";")]
                try:
                    for i, member in enumerate(members):
                        replace_placeholder_in_cell(cell, f"{{f{8 + 2 * i}}}", member[0])
                        replace_placeholder_in_cell(cell, f"{{f{9 + 2 * i}}}", member[1])
                except:
                    pass

    # 나머지 자리표시자 초기화
    for i in range(8, 16):
        placeholder = f"{{f{i}}}"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholder_in_cell(cell, placeholder, "")

    # 문서 저장
    output_path = f"doc_conv/회의록_{doc_s_no}.docx"
    doc.save(output_path)

    push.push_to_nextjs(output_path, f"회의록_{doc_s_no}.docx")

    return {"RESULT_CODE": 200, "RESULT_MSG": "Done!"}

def process_summary(doc_s_no):
    """
    프로젝트 개요서 문서를 생성하는 기능을 처리합니다.
    """
    # DB에서 프로젝트 개요서 데이터 조회
    summary_data = output_DB.fetch_one_summary_document(doc_s_no)
    if not summary_data:
        raise HTTPException(status_code=404, detail="Project summary not found")

    try:
        # 템플릿 로드
        doc = Document("/data/Docs_Template/개요서.docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template loading error: {e}")

    # 자리표시자 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_in_cell(cell, "{f1}", summary_data.get("doc_s_name", ""))
                replace_placeholder_in_cell(
                    cell, "{f2}",
                    summary_data.get("doc_s_start", "").strftime("%Y-%m-%d")
                    if isinstance(summary_data.get("doc_s_start"), date)
                    else str(summary_data.get("doc_s_start", ""))
                )
                replace_placeholder_in_cell(cell, "{f3}", summary_data.get("doc_s_team", ""))
                replace_placeholder_in_cell(
                    cell, "{f4}",
                    summary_data.get("doc_s_end", "").strftime("%Y-%m-%d")
                    if isinstance(summary_data.get("doc_s_end"), date)
                    else str(summary_data.get("doc_s_end", ""))
                )
                replace_placeholder_in_cell(
                    cell, "{f5}",
                    summary_data.get("doc_s_date", "").strftime("%Y-%m-%d")
                    if isinstance(summary_data.get("doc_s_date"), date)
                    else str(summary_data.get("doc_s_date", ""))
                )
                replace_placeholder_in_cell(cell, "{f6}", summary_data.get("doc_s_overview", ""))
                replace_placeholder_in_cell(cell, "{f7}", summary_data.get("doc_s_goals", ""))
                replace_placeholder_in_cell(cell, "{f8}", summary_data.get("doc_s_range", ""))
                replace_placeholder_in_cell(cell, "{f9}", summary_data.get("doc_s_stack", ""))
                replace_placeholder_in_cell(cell, "{f10}", summary_data.get("doc_s_outcomes", ""))

    # 문서 저장
    output_path = f"doc_conv/개요서_{doc_s_no}.docx"
    doc.save(output_path)

    push.push_to_nextjs(output_path, f"개요서_{doc_s_no}.docx")

    return {"RESULT_CODE": 200, "RESULT_MSG": "Done!"}


def process_reqspec(doc_r_no):
    """
    요구사항 명세서 문서를 생성하는 기능을 처리합니다.
    """
    # DB에서 요구사항 명세서 데이터 조회
    reqspec_data = output_DB.fetch_one_reqspec(doc_r_no)
    if not reqspec_data:
        raise HTTPException(status_code=404, detail="Requirement specification not found")

    try:
        # 템플릿 로드
        doc = Document("/data/Docs_Template/요구사항.docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template loading error: {e}")

    # 자리표시자 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_in_cell(
                    cell, "{f1}",
                    reqspec_data.get("doc_r_date", "").strftime("%Y-%m-%d")
                    if isinstance(reqspec_data.get("doc_r_date"), date)
                    else str(reqspec_data.get("doc_r_date", ""))
                )
                replace_placeholder_in_cell(cell, "{f2}", reqspec_data.get("doc_r_s_name", ""))
                replace_placeholder_in_cell(cell, "{f3}", reqspec_data.get("doc_r_s_content", ""))
                replace_placeholder_in_cell(cell, "{f4}", reqspec_data.get("doc_r_f_name", ""))
                replace_placeholder_in_cell(cell, "{f5}", reqspec_data.get("doc_r_f_content", ""))
                replace_placeholder_in_cell(cell, "{f6}", str(reqspec_data.get("doc_r_f_priority", "")))
                replace_placeholder_in_cell(cell, "{f7}", reqspec_data.get("doc_r_nf_name", ""))
                replace_placeholder_in_cell(cell, "{f8}", reqspec_data.get("doc_r_nf_content", ""))
                replace_placeholder_in_cell(cell, "{f9}", str(reqspec_data.get("doc_r_nf_priority", "")))

    # for paragraph in doc.paragraphs:
    #     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 문서 저장
    output_path = f"doc_conv/요구사항_{doc_r_no}.docx"
    doc.save(output_path)

    push.push_to_nextjs(output_path, f"요구사항_{doc_r_no}.docx")

    return {"RESULT_CODE": 200, "RESULT_MSG": "Done!"}


def process_testcase(doc_t_no):
    """
    특정 테스트 케이스 문서를 생성하는 기능을 처리합니다.
    """
    # DB에서 특정 테스트 케이스 데이터 조회
    testcase_data = output_DB.fetch_one_testcase(doc_t_no)
    if not testcase_data:
        raise HTTPException(status_code=404, detail="Test case not found")

    try:
        # 템플릿 로드
        doc = Document("/data/Docs_Template/테스트케이스.docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template loading error: {e}")

    # 첫 번째 항목 기준으로 자리표시자 치환
    first_case = testcase_data[0] if testcase_data else {}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_in_cell(
                    cell, "{f1}",
                    first_case.get("doc_t_start", "").strftime("%Y-%m-%d")
                    if isinstance(first_case.get("doc_t_start"), date)
                    else str(first_case.get("doc_t_start", ""))
                )
                replace_placeholder_in_cell(
                    cell, "{f2}",
                    first_case.get("doc_t_end", "").strftime("%Y-%m-%d")
                    if isinstance(first_case.get("doc_t_end"), date)
                    else str(first_case.get("doc_t_end", ""))
                )
                replace_placeholder_in_cell(
                    cell, "{f3}",
                    str(first_case.get("doc_t_name", ""))
                )
                replace_placeholder_in_cell(
                    cell, "{f4}",
                    str(first_case.get("doc_t_pass", ""))
                )

    # 문서 저장
    output_path = f"doc_conv/테스트케이스_{doc_t_no}.docx"
    doc.save(output_path)

    push.push_to_nextjs(output_path, f"테스트케이스_{doc_t_no}.docx")

    return {"RESULT_CODE": 200, "RESULT_MSG": "Done!", "OUTPUT_PATH": output_path}


def process_report(doc_rep_no):
    """
    보고서 문서를 생성하는 기능을 처리합니다.
    """
    # DB에서 보고서 데이터 조회
    report_data = output_DB.fetch_one_report(doc_rep_no)
    if not report_data:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        # 템플릿 로드
        doc = Document("/data/Docs_Template/보고서.docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template loading error: {e}")

    # 자리표시자 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_in_cell(cell, "{f1}", report_data.get("doc_rep_name", ""))
                replace_placeholder_in_cell(cell, "{f2}", report_data.get("doc_rep_pname", ""))
                replace_placeholder_in_cell(
                    cell, "{f3}",
                    report_data.get("doc_rep_date", "").strftime("%Y-%m-%d")
                    if isinstance(report_data.get("doc_rep_date"), date)
                    else str(report_data.get("doc_rep_date", ""))
                )
                replace_placeholder_in_cell(cell, "{f8}", report_data.get("doc_rep_professor", ""))
                replace_placeholder_in_cell(cell, "{f9}", report_data.get("doc_rep_writer", ""))
                replace_placeholder_in_cell(cell, "{f10}", report_data.get("doc_rep_research", ""))
                replace_placeholder_in_cell(cell, "{f11}", report_data.get("doc_rep_design", ""))
                replace_placeholder_in_cell(cell, "{f12}", report_data.get("doc_rep_arch", ""))
                replace_placeholder_in_cell(cell, "{f13}", report_data.get("doc_rep_result", ""))
                replace_placeholder_in_cell(cell, "{f14}", report_data.get("doc_rep_conclusion", ""))

                # 팀원 목록 치환
                members = [member.split(",")[0] for member in report_data.get("doc_rep_member", "").split(";")]
                for i, member in enumerate(members[:4]):
                    replace_placeholder_in_cell(cell, f"{{f{4 + i}}}", member)

                # 남은 자리표시자 제거
                for i in range(1, 15):
                    replace_placeholder_in_cell(cell, f"{{f{i}}}", "")

    # 문서 저장
    output_path = f"doc_conv/보고서_{doc_rep_no}.docx"
    doc.save(output_path)

    push.push_to_nextjs(output_path, f"보고서_{doc_rep_no}.docx")

    return {"RESULT_CODE": 200, "RESULT_MSG": "Done!", "OUTPUT_PATH": output_path}


@router.post("/docs/convert")
async def docs_convert(payload: ConverterPayload):
    try:
        if payload.doc_type == 0: # 프로젝트 개요서
            return process_summary(payload.doc_s_no)
        elif payload.doc_type == 1:  # 회의록
            return process_meeting_minutes(payload.doc_s_no)
        elif payload.doc_type == 2: # 테스트 케이스
            # return process_testcase(payload.doc_s_no) # 테스트 케이스 컨셉 변경에 따라 문서 변환 기능 비활성화 (25.03.23)
            return {"RESULT_CODE": 410, "RESULT_MSG": "테스트 케이스의 문서 변환 기능은 지원이 종료됐습니다."}
        elif payload.doc_type == 3: # 요구사항 명세서
            return process_reqspec(payload.doc_s_no)
        elif payload.doc_type == 4: # 보고서
            return process_report(payload.doc_s_no)
        else:
            raise HTTPException(status_code=400, detail="Unsupported document type")
    except Exception:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Unexpected error occurred")
